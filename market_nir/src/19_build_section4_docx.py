#!/usr/bin/env python3
from __future__ import annotations

import argparse
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

SECTION_TITLE = (
    "4 ВЛИЯНИЕ АРХИТЕКТУРЫ МОДЕЛИ НА КАЧЕСТВО И ВЫЧИСЛИТЕЛЬНУЮ "
    "ЭФФЕКТИВНОСТЬ ПРОГНОЗИРОВАНИЯ РЫНОЧНЫХ ВРЕМЕННЫХ РЯДОВ"
)
OLD_SECTION_TITLE = "4 СВОДНЫЙ СРАВНИТЕЛЬНЫЙ АНАЛИЗ"


def parse_args() -> argparse.Namespace:
    p = argparse.ArgumentParser(description="Build detailed course DOCX with section 4 market architecture analysis")
    p.add_argument("--input-docx", default="Курсовая_работа_Савченко_КА.docx")
    p.add_argument("--output-docx", default="Курсовая_работа_Савченко_КА_block4_market_detailed.docx")
    p.add_argument("--artifacts", default="market_nir/artifacts/architecture_comparison")
    return p.parse_args()


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, size: int = 8, align=WD_ALIGN_PARAGRAPH.CENTER) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(str(text))
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def insert_paragraph_before(anchor, text: str = "", style: str | None = None, bold: bool = False):
    p = anchor.insert_paragraph_before(text)
    if style:
        p.style = style
    for run in p.runs:
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)
        run.bold = bold or run.bold
    return p


def insert_heading(anchor, text: str, level: int):
    style = "Heading 1" if level == 1 else "Heading 2" if level == 2 else "Heading 3"
    p = insert_paragraph_before(anchor, text, style=style, bold=True)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0, 0, 0)
        run.font.size = Pt(14 if level == 1 else 13 if level == 2 else 12)
    return p


def insert_picture(anchor, path: Path, caption: str, width: float = 6.35) -> None:
    p = anchor.insert_paragraph_before()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width))
    cap = anchor.insert_paragraph_before(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.name = "Times New Roman"
        cap.runs[0].font.size = Pt(10)


def insert_table(anchor, rows: list[list[str]], caption: str, font_size: int = 8) -> None:
    doc = anchor._parent
    table = doc.add_table(rows=len(rows), cols=len(rows[0]), width=Inches(6.4))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            align = WD_ALIGN_PARAGRAPH.LEFT if j == 0 and i > 0 else WD_ALIGN_PARAGRAPH.CENTER
            set_cell_text(table.cell(i, j), str(val), bold=(i == 0), size=font_size, align=align)
            if i == 0:
                set_cell_shading(table.cell(i, j), "D9EAF7")
            elif i % 2 == 0:
                set_cell_shading(table.cell(i, j), "F4F8FB")
    tbl = table._tbl
    tbl.getparent().remove(tbl)
    anchor._element.addprevious(tbl)
    cap = anchor.insert_paragraph_before(caption)
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if cap.runs:
        cap.runs[0].italic = True
        cap.runs[0].font.size = Pt(10)


def remove_body_section4(doc: Document):
    paras = doc.paragraphs
    starts = [i for i, p in enumerate(paras) if p.text.strip().startswith(OLD_SECTION_TITLE)]
    if not starts:
        raise RuntimeError("Old section 4 was not found")
    start = max(starts)  # skip old TOC entry, use body occurrence
    ends = [i for i, p in enumerate(paras) if i > start and p.text.strip() == "ЗАКЛЮЧЕНИЕ"]
    if not ends:
        raise RuntimeError("Conclusion after old section 4 was not found")
    end = ends[0]
    anchor = paras[end]
    for p in paras[start:end]:
        el = p._element
        el.getparent().remove(el)
    return anchor


def patch_toc_and_conclusion(doc: Document) -> None:
    replacements = {
        OLD_SECTION_TITLE: SECTION_TITLE,
        "Выполнена расширенная научно-исследовательская переработка курсового отчета по трем направлениям: Transformer, GAN и GNN.":
            "Выполнена расширенная научно-исследовательская переработка курсового отчета по четырем направлениям: Transformer, GAN, GNN и сравнительный анализ архитектур для прогнозирования рыночных временных рядов.",
        "Ключевой результат текущей версии - переход от демонстрационного формата к НИР-представлению: текстовая часть существенно углублена, а визуальная составляющая расширена до 35+ аналитических иллюстраций, встроенных непосредственно в контекст соответствующих выводов.":
            "Ключевой результат текущей версии - переход от демонстрационного формата к НИР-представлению: текстовая часть существенно углублена, а визуальная составляющая расширена до 50+ аналитических иллюстраций, встроенных непосредственно в контекст соответствующих выводов.",
        "Всего иллюстраций в отчете: 35.": "Всего иллюстраций в отчете после добавления расширенного рыночного раздела: 60.",
        "Итоговое число рисунков после расширения: 41.": "Итоговое число рисунков после расширения и добавления рыночного раздела: 60.",
    }
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt in replacements and p.runs:
            for run in p.runs:
                run.text = ""
            p.runs[0].text = replacements[txt]


def metric_rows(metrics: pd.DataFrame) -> list[list[str]]:
    rows = [["Модель", "Bal.acc", "Macro-F1", "Hit-rate top-10%", "PnL top-10%", "Train, c", "мс/1000", "Параметры"]]
    for _, r in metrics.iterrows():
        rows.append([
            r["model"],
            f"{r['balanced_accuracy']:.3f}",
            f"{r['macro_f1']:.3f}",
            f"{r['top10_hit_rate']:.3f}",
            f"{r['top10_cum_return']:.3f}",
            f"{r['train_seconds']:.2f}",
            f"{r['inference_ms_per_1000']:.2f}",
            f"{int(r['params_count'])}",
        ])
    return rows


def main() -> None:
    args = parse_args()
    root = Path.cwd()
    artifacts = root / args.artifacts
    plots = artifacts / "plots"
    metrics = pd.read_csv(artifacts / "metrics" / "architecture_comparison_metrics.csv")
    histories = pd.read_json(artifacts / "metrics" / "training_histories.json") if (artifacts / "metrics" / "training_histories.json").exists() else None

    doc = Document(root / args.input_docx)
    anchor = remove_body_section4(doc)
    patch_toc_and_conclusion(doc)

    best = metrics.iloc[0]
    fastest = metrics.sort_values("train_seconds").iloc[0]

    insert_heading(anchor, SECTION_TITLE, 1)
    insert_heading(anchor, "4.1 Исследовательская постановка и границы задачи", 2)
    insert_paragraph_before(anchor, "В четвертом разделе исследуется не задача построения готовой торговой системы, а влияние выбора архитектуры модели на качество и вычислительную эффективность прогнозирования рыночных временных рядов. Такая формулировка важна, потому что рыночные данные принципиально отличаются от учебных синтетических наборов: полезный сигнал слабый, распределения меняются во времени, а итоговая метрика может быстро деградировать при переносе на следующий период.")
    insert_paragraph_before(anchor, "Практическая цель раздела состоит в том, чтобы получить воспроизводимый экспериментальный стенд: одни и те же OHLCV-данные, один и тот же способ формирования целевой переменной, единое временное разбиение и несколько архитектур разной природы. В этом случае различия в метриках можно обсуждать как следствие архитектурных предположений модели, а не как результат случайной подготовки данных.")
    insert_paragraph_before(anchor, "Гипотеза H4.1: модели, способные учитывать нелинейные зависимости и короткий временной контекст, должны превзойти линейный baseline. Гипотеза H4.2: повышение качества будет сопровождаться ростом вычислительной стоимости. Гипотеза H4.3: классификационная метрика сама по себе недостаточна, поэтому дополнительно требуется проверять связь score модели с будущей доходностью.")
    insert_picture(anchor, plots / "4_00_market_experiment_pipeline.png", "Рисунок 4.1 - Pipeline эксперимента раздела 4")

    insert_heading(anchor, "4.2 Рыночные данные и формирование целевой переменной", 2)
    insert_paragraph_before(anchor, "Исходными данными являются рыночные бары OHLCV: open, high, low, close и volume для каждого тикера и момента времени. В текущем локальном прогоне использован датасет из 15039 объектов, двух тикеров и 164 числовых признаков. Для каждого объекта вычисляется будущая доходность ret_h на заданном горизонте. Затем по динамическому порогу, связанному с локальной волатильностью, объект получает метку UP или DOWN.")
    insert_paragraph_before(anchor, "Важно, что разбиение выполняется по времени, а не случайно. Это защищает эксперимент от утечки будущего в прошлое: модель обучается на раннем участке, валидируется на следующем и тестируется на самом позднем фрагменте. Для рыночных задач такой протокол принципиален, потому что случайный train/test split обычно завышает качество за счет перемешивания близких по времени наблюдений.")
    insert_picture(anchor, plots / "4_09_split_and_label_distribution.png", "Рисунок 4.2 - Размеры train/val/test и распределение классов")
    insert_paragraph_before(anchor, "Рисунок 4.2 показывает, что классы UP и DOWN в целом сопоставимы по количеству, поэтому задача не сводится к угадыванию доминирующего класса. Одновременно распределение классов различается между временными отрезками, что отражает нестационарность рынка и объясняет, почему balanced accuracy используется как одна из основных метрик.")
    insert_picture(anchor, plots / "4_11_market_return_context.png", "Рисунок 4.3 - Накопленная будущая доходность ret_h по тикерам")

    insert_heading(anchor, "4.3 Конструирование признаков", 2)
    insert_paragraph_before(anchor, "Признаки строятся только из рыночных данных. В базовый блок входят лаговые доходности, rolling-средние и rolling-стандартные отклонения, признаки объема, RSI, ATR, MACD, Bollinger z-score, календарные признаки часа и дня недели. Дополнительно добавлен market-context: средние движения остальных активов, relative-to-market признаки, benchmark-relative признаки и market breadth. Последняя группа позволяет модели видеть не только состояние конкретного тикера, но и фон рынка.")
    insert_paragraph_before(anchor, "С точки зрения честности эксперимента важно, что все признаки вычисляются из прошлого и текущего бара. Будущая доходность используется только как целевая переменная. Поэтому модель не получает прямого доступа к информации, которую должна предсказать.")
    insert_picture(anchor, plots / "4_10_feature_group_counts.png", "Рисунок 4.4 - Группы признаков market-only датасета")

    insert_heading(anchor, "4.4 Обоснование выбора архитектур", 2)
    insert_paragraph_before(anchor, "Для сравнения выбраны пять архитектур, различающихся индуктивными предположениями. Logistic Regression задает нижнюю границу сложности: если она показывает результат около случайного уровня, значит линейной комбинации признаков недостаточно. HistGradientBoosting проверяет, помогают ли нелинейные табличные взаимодействия. TorchMLP добавляет нейросетевую нелинейность без явной временной памяти. TemporalCNN проверяет гипотезу о локальных паттернах в коротком окне. TinyTransformer проверяет, полезен ли механизм self-attention для адаптивного учета зависимостей между шагами окна.")
    rows = [
        ["Архитектура", "Вход", "Индуктивное предположение", "Ожидаемый риск"],
        ["LogisticRegression", "Один вектор признаков", "Линейная разделимость", "Недообучение"],
        ["HistGradientBoosting", "Один вектор признаков", "Нелинейные табличные правила", "Переобучение на шум"],
        ["TorchMLP", "Один вектор признаков", "Нелинейная комбинация признаков", "Чувствительность к масштабу"],
        ["TemporalCNN", "Окно из 12 шагов", "Локальные временные мотивы", "Слабость к дальним связям"],
        ["TinyTransformer", "Окно из 12 шагов", "Self-attention между шагами", "Больше параметров и время"],
    ]
    insert_table(anchor, rows, "Таблица 4.1 - Архитектуры и их исследовательские предположения", font_size=8)
    insert_picture(anchor, plots / "4_12_architecture_assumptions.png", "Рисунок 4.5 - Сравниваемые архитектуры и их индуктивные предположения")

    insert_heading(anchor, "4.5 Подробный процесс обучения", 2)
    insert_paragraph_before(anchor, "Для классических моделей обучение выполняется как оптимизация на фиксированной матрице признаков. Logistic Regression предварительно стандартизует признаки и минимизирует регуляризованную логистическую потерю. HistGradientBoosting строит ансамбль деревьев последовательно: каждое новое дерево исправляет ошибки предыдущих, а глубина и минимальное число объектов в листе ограничивают переобучение.")
    insert_paragraph_before(anchor, "Для нейросетевых моделей используется PyTorch. Перед обучением признаки стандартизируются по train-части, затем те же параметры масштаба применяются к val/test. Оптимизатор AdamW обновляет веса по mini-batch, функция потерь - class-weighted CrossEntropyLoss, что уменьшает влияние дисбаланса классов. На каждом шаге применяется gradient clipping, чтобы предотвратить нестабильные скачки градиента. После каждой эпохи считается balanced accuracy на validation-части; сохраняется состояние модели с лучшим validation-результатом.")
    rows = [
        ["Модель", "Оптимизация", "Регуляризация", "Validation-критерий"],
        ["LogisticRegression", "LBFGS", "class_weight=balanced", "test после fit"],
        ["HistGradientBoosting", "градиентный бустинг", "depth, min_samples_leaf, l2", "test после fit"],
        ["TorchMLP", "AdamW", "dropout, weight decay, clipping", "val balanced accuracy"],
        ["TemporalCNN", "AdamW", "weight decay, clipping", "val balanced accuracy"],
        ["TinyTransformer", "AdamW", "dropout, weight decay, clipping", "val balanced accuracy"],
    ]
    insert_table(anchor, rows, "Таблица 4.2 - Протокол обучения моделей", font_size=8)
    insert_picture(anchor, plots / "4_13_training_val_balanced_accuracy.png", "Рисунок 4.6 - Validation balanced accuracy нейросетевых моделей по эпохам")
    insert_paragraph_before(anchor, "Кривая validation balanced accuracy показывает, что снижение train loss не всегда означает улучшение обобщающей способности. Например, нейросеть может продолжать уменьшать ошибку на обучающей выборке, но validation-метрика перестает расти. Поэтому в эксперименте используется ранняя остановка по validation-качеству, а не выбор последней эпохи.")
    insert_picture(anchor, plots / "4_14_training_loss_curves.png", "Рисунок 4.7 - Training loss нейросетевых моделей")

    insert_heading(anchor, "4.6 Метрики качества и вычислительной эффективности", 2)
    insert_paragraph_before(anchor, "Accuracy показывает долю правильных ответов, но для рыночных данных она недостаточна: если один класс временно преобладает, модель может выглядеть лучше, чем есть. Balanced accuracy усредняет recall по классам и устойчивее к дисбалансу. Macro-F1 дополнительно учитывает precision и recall каждого класса. Score-ret correlation показывает, согласован ли непрерывный score модели с будущей доходностью. Top-10% hit-rate и PnL проверяют только наиболее уверенные предсказания: это не торговая стратегия, а диагностический тест полезности score.")
    insert_table(anchor, metric_rows(metrics), "Таблица 4.3 - Итоговые метрики архитектур", font_size=7)
    insert_picture(anchor, plots / "4_01_architecture_quality_metrics.png", "Рисунок 4.8 - Сравнение архитектур по balanced accuracy, macro-F1 и hit-rate")
    insert_paragraph_before(anchor, f"Лучший результат в текущем прогоне показала модель {best['model']}: balanced accuracy = {best['balanced_accuracy']:.3f}, macro-F1 = {best['macro_f1']:.3f}. При этом самая быстрая модель по обучению - {fastest['model']}, что подчеркивает различие между качеством и вычислительной стоимостью.")
    insert_picture(anchor, plots / "4_08_metrics_heatmap.png", "Рисунок 4.9 - Тепловая карта итоговых метрик")

    insert_heading(anchor, "4.7 Анализ результатов по архитектурам", 2)
    insert_paragraph_before(anchor, "Logistic Regression оказалась полезной как baseline: она обучается почти мгновенно и имеет минимальное число параметров, но качество близко к границе случайного угадывания. HistGradientBoosting в данном запуске не превзошел линейную модель, что можно объяснить высокой шумностью данных и ограниченной длиной истории. TorchMLP получил более высокую balanced accuracy, но его top-signal диагностика осталась отрицательной по PnL, значит классификационная уверенность не полностью совпала с направлением доходности.")
    insert_paragraph_before(anchor, "TemporalCNN использует локальные свертки по временному окну, однако на выбранном наборе признаков локальные шаблоны оказались недостаточно устойчивыми. TinyTransformer показал лучший баланс классификационного качества и направленного score: механизм внимания позволяет модели по-разному взвешивать шаги внутри окна, а не применять один и тот же локальный фильтр ко всем позициям.")
    insert_picture(anchor, plots / "4_02_quality_vs_train_time.png", "Рисунок 4.10 - Компромисс качества и времени обучения")
    insert_picture(anchor, plots / "4_03_inference_latency.png", "Рисунок 4.11 - Скорость инференса разных архитектур")
    insert_picture(anchor, plots / "4_04_parameter_complexity.png", "Рисунок 4.12 - Параметрическая сложность моделей")
    insert_picture(anchor, plots / "4_18_quality_efficiency_radar.png", "Рисунок 4.13 - Интегральный профиль качества и эффективности")

    insert_heading(anchor, "4.8 Диагностика score и прикладная интерпретация", 2)
    insert_paragraph_before(anchor, "Для рыночной задачи важно анализировать не только дискретный класс, но и непрерывный score = P(UP) - P(DOWN). Если score положительно связан с будущей доходностью, модель несет хотя бы слабую направленную информацию. Если такой связи нет, высокая accuracy может быть случайной или не пригодной для практической интерпретации.")
    insert_picture(anchor, plots / "4_15_score_distributions.png", "Рисунок 4.14 - Распределение score по моделям")
    insert_picture(anchor, plots / "4_06_score_vs_future_return.png", "Рисунок 4.15 - Связь score лучшей модели и будущей доходности")
    insert_paragraph_before(anchor, "Наиболее осторожная интерпретация состоит в следующем: положительная связь score и ret_h не означает готовую стратегию, но показывает, что модель не просто угадывает классы формально. Поэтому top-10% сигналов используются как дополнительный стресс-тест: берутся только самые уверенные предсказания, после чего проверяется знак фактической доходности.")
    insert_picture(anchor, plots / "4_16_top10_pnl_by_model.png", "Рисунок 4.16 - PnL top-10% наиболее уверенных сигналов")
    insert_picture(anchor, plots / "4_07_equity_curves_top10.png", "Рисунок 4.17 - Equity-кривые top-10% сигналов")

    insert_heading(anchor, "4.9 Ошибки, устойчивость и ограничения", 2)
    insert_paragraph_before(anchor, "Матрица ошибок показывает, насколько симметрично модель путает UP и DOWN. Для рыночных данных полностью устранить такие ошибки невозможно: часть движений действительно близка к шуму, а часть возникает под воздействием внешних событий, отсутствующих в OHLCV. Поэтому отдельный интерес представляет не только общий процент ошибок, но и их поведение во времени.")
    insert_picture(anchor, plots / "4_05_best_model_confusion_matrix.png", "Рисунок 4.18 - Матрица ошибок лучшей модели")
    insert_picture(anchor, plots / "4_17_best_model_rolling_error.png", "Рисунок 4.19 - Скользящая доля ошибок лучшей модели")
    insert_paragraph_before(anchor, "Основные ограничения эксперимента: небольшой набор тикеров в локальном прогоне, ограниченный исторический период, отсутствие новостных и макроэкономических факторов, а также чувствительность результата к горизонту прогнозирования и способу отбора top-сигналов. Эти ограничения не обесценивают раздел, но задают корректные границы вывода: исследуется влияние архитектуры в контролируемой постановке, а не строится универсальная торговая система.")

    insert_heading(anchor, "4.10 Выводы по разделу", 2)
    insert_paragraph_before(anchor, "Раздел показывает, что архитектура модели влияет не только на итоговую точность, но и на вычислительную стоимость, характер ошибок и полезность непрерывного score. В текущем прогоне TinyTransformer оказался лучшим по balanced accuracy и macro-F1, а также дал наиболее убедительную top-signal диагностику. Однако его преимущество сопровождается большим числом параметров и более высокой стоимостью обучения по сравнению с линейной моделью.")
    insert_paragraph_before(anchor, "Главный вывод: для рыночных временных рядов нельзя выбирать модель только по названию архитектуры. Необходимо одновременно анализировать качество, устойчивость на validation/test, вычислительный бюджет, интерпретируемость score и поведение ошибок во времени. Именно такой многокритериальный подход делает эксперимент пригодным для курсовой работы: результат не выглядит искусственно идеальным, но дает содержательную основу для сравнения архитектур.")

    for p in doc.paragraphs:
        for run in p.runs:
            if run.font.name is None:
                run.font.name = "Times New Roman"
            if run.font.size is None:
                run.font.size = Pt(12)

    doc.save(root / args.output_docx)
    print(root / args.output_docx)


if __name__ == "__main__":
    main()
